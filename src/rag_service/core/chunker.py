"""Document chunking and processing using unstructured library."""

import logging
import os
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from typing import Literal

import hashlib

from rag_service.core.exceptions import DocumentProcessingError, UnsupportedFileTypeError
from rag_service.core.retriever import Document

logger = logging.getLogger(__name__)


# Supported file extensions
SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".txt",
    ".md",
    ".markdown",
    ".xml",
    ".html",
    ".htm",
    ".docx",
    ".doc",
    ".py",
    ".js",
    ".ts",
    ".jsx",
    ".tsx",
    ".java",
    ".c",
    ".cpp",
    ".h",
    ".hpp",
    ".rs",
    ".go",
    ".rb",
    ".php",
    ".json",
    ".yaml",
    ".yml",
    ".toml",
    ".ini",
    ".cfg",
    ".conf",
    ".sh",
    ".bash",
    ".zsh",
    ".ps1",
    ".bat",
    ".cmd",
}


class DocumentChunker:
    """Document chunker that processes various file types and splits into chunks."""

    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 50,
        pdf_strategy: Literal["fast", "hi_res", "ocr_only"] = "fast",
    ) -> None:
        """Initialize the document chunker.

        Args:
            chunk_size: Target size of chunks in characters.
            chunk_overlap: Number of characters to overlap between chunks.
            pdf_strategy: Strategy for PDF processing.
        """
        self._chunk_size = chunk_size
        self._chunk_overlap = chunk_overlap
        self._pdf_strategy = pdf_strategy

    def process_file(self, file_path: Path | str) -> list[Document]:
        """Process a single file and return documents.

        Args:
            file_path: Path to the file to process.

        Returns:
            List of Document objects with chunked text.

        Raises:
            UnsupportedFileTypeError: If file type is not supported.
            DocumentProcessingError: If processing fails.
        """
        path = Path(file_path)

        if not path.exists():
            raise DocumentProcessingError(f"File not found: {path}")

        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_EXTENSIONS:
            raise UnsupportedFileTypeError(
                f"Unsupported file type: {suffix}. Supported: {SUPPORTED_EXTENSIONS}"
            )

        try:
            # Extract text from file
            text = self._extract_text(path)

            # Chunk the text
            chunks = self._chunk_text(text)

            # Create documents
            documents = []
            for i, chunk in enumerate(chunks):
                doc_id = self._generate_doc_id(path, i)
                documents.append(
                    Document(
                        text=chunk,
                        metadata={
                            "source": str(path),
                            "filename": path.name,
                            "chunk_index": i,
                            "total_chunks": len(chunks),
                            "file_type": suffix,
                        },
                        document_id=doc_id,
                    )
                )

            return documents

        except Exception as e:
            raise DocumentProcessingError(f"Failed to process {path}: {e}") from e

    def process_directory(
        self,
        directory: Path | str,
        recursive: bool = True,
        max_workers: int | None = None,
    ) -> list[Document]:
        """Process all supported files in a directory using parallel processing.

        Processes files in parallel using ThreadPoolExecutor to utilize all
        available CPU cores. This significantly speeds up processing of
        multiple files.

        Args:
            directory: Path to the directory to process.
            recursive: Whether to process subdirectories.
            max_workers: Maximum number of worker threads. Defaults to CPU count.

        Returns:
            List of Document objects from all files.
        """
        path = Path(directory)

        if not path.exists():
            raise DocumentProcessingError(f"Directory not found: {path}")

        if not path.is_dir():
            raise DocumentProcessingError(f"Not a directory: {path}")

        # Collect all file paths first
        pattern = "**/*" if recursive else "*"
        file_paths = [
            file_path
            for file_path in path.glob(pattern)
            if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_EXTENSIONS
        ]

        if not file_paths:
            logger.info(f"No supported files found in {path}")
            return []

        # Determine optimal worker count
        if max_workers is None:
            cpu_count = os.cpu_count() or 4
            max_workers = min(cpu_count, len(file_paths))
            # Cap at reasonable limit to avoid excessive thread overhead
            max_workers = min(max_workers, 16)

        logger.info(
            f"Processing {len(file_paths)} files in parallel using {max_workers} workers"
        )

        documents = []
        errors = []

        # Process files in parallel
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Submit all file processing tasks
            future_to_path = {
                executor.submit(self.process_file, file_path): file_path
                for file_path in file_paths
            }

            # Collect results as they complete
            for future in as_completed(future_to_path):
                file_path = future_to_path[future]
                try:
                    file_docs = future.result()
                    documents.extend(file_docs)
                    logger.debug(f"Processed {file_path.name}: {len(file_docs)} chunks")
                except (UnsupportedFileTypeError, DocumentProcessingError) as e:
                    logger.warning(f"Skipping {file_path.name}: {e}")
                    errors.append(str(file_path))
                    continue
                except Exception as e:
                    logger.error(f"Error processing {file_path.name}: {e}")
                    errors.append(str(file_path))
                    continue

        if errors:
            logger.warning(f"Failed to process {len(errors)} files: {errors[:5]}...")

        logger.info(
            f"Completed processing: {len(documents)} total chunks from {len(file_paths) - len(errors)} files"
        )

        return documents

    def _extract_text(self, path: Path) -> str:
        """Extract text content from a file.

        Args:
            path: Path to the file.

        Returns:
            Extracted text content.
        """
        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return self._extract_pdf(path)
        elif suffix in {".txt", ".md", ".markdown"}:
            return self._extract_text_file(path)
        elif suffix == ".xml":
            return self._extract_xml(path)
        elif suffix in {".html", ".htm"}:
            return self._extract_html(path)
        elif suffix == ".docx":
            return self._extract_docx(path)
        else:
            # Treat as plain text (code files, config files, etc.)
            return self._extract_text_file(path)

    def _extract_pdf(self, path: Path) -> str:
        """Extract text from PDF using unstructured or pypdf."""
        try:
            # Try unstructured first for better results
            from unstructured.partition.pdf import partition_pdf

            elements = partition_pdf(
                str(path),
                strategy=self._pdf_strategy,
                extract_images_in_pdf=False,
            )
            return "\n\n".join(el.text for el in elements if el.text.strip())
        except (ImportError, Exception) as e:
            # Fallback to pypdf if unstructured fails (missing poppler, etc.)
            error_msg = str(e).lower()
            if "poppler" in error_msg or "page count" in error_msg:
                logger.warning(
                    f"Poppler not available for PDF processing. Falling back to pypdf for {path.name}. "
                    "Install Poppler for better PDF extraction: https://github.com/oschwartz10612/poppler-windows/releases"
                )
            else:
                logger.warning(f"unstructured PDF processing failed for {path.name}: {e}. Falling back to pypdf.")
            
            # Fallback to pypdf
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n\n".join(text_parts)

    def _extract_text_file(self, path: Path) -> str:
        """Extract text from plain text files."""
        encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
        for encoding in encodings:
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        raise DocumentProcessingError(f"Could not decode file: {path}")

    def _extract_xml(self, path: Path) -> str:
        """Extract text from XML files."""
        try:
            from unstructured.partition.xml import partition_xml

            elements = partition_xml(str(path))
            return "\n\n".join(el.text for el in elements if el.text.strip())
        except ImportError:
            # Fallback to plain text extraction
            return self._extract_text_file(path)

    def _extract_html(self, path: Path) -> str:
        """Extract text from HTML files."""
        try:
            from unstructured.partition.html import partition_html

            elements = partition_html(str(path))
            return "\n\n".join(el.text for el in elements if el.text.strip())
        except ImportError:
            # Fallback to plain text extraction
            return self._extract_text_file(path)

    def _extract_docx(self, path: Path) -> str:
        """Extract text from DOCX files.

        Uses python-docx directly for speed. The unstructured library's
        partition_docx is slower due to table/image analysis overhead.

        Optimized with list comprehensions for better performance.
        """
        logger.debug(f"Extracting text from DOCX: {path.name}")

        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(str(path))

            # Extract paragraphs (optimized)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

            # Extract tables (optimized with list comprehension)
            table_texts = [
                " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                for table in doc.tables
                for row in table.rows
            ]
            paragraphs.extend(text for text in table_texts if text)

            text = "\n\n".join(paragraphs)
            logger.debug(f"Extracted {len(text):,} characters from {path.name}")
            return text

        except ImportError:
            raise DocumentProcessingError(
                "DOCX processing requires 'python-docx' package. "
                "Install it with: pip install python-docx"
            )
        except Exception as e:
            raise DocumentProcessingError(f"Failed to process DOCX file {path.name}: {e}")

    def _chunk_text(self, text: str) -> list[str]:
        """Split text into overlapping chunks.

        Uses a simple character-based splitting with overlap.

        Args:
            text: Text to split.

        Returns:
            List of text chunks.
        """
        if not text.strip():
            return []

        # Normalize whitespace
        text = " ".join(text.split())

        if len(text) <= self._chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            # Find end position
            end = start + self._chunk_size

            if end >= len(text):
                # Last chunk
                chunks.append(text[start:].strip())
                break

            # Try to break at a sentence boundary
            chunk_text = text[start:end]
            break_chars = [". ", "! ", "? ", "\n", "; "]

            best_break = -1
            for char in break_chars:
                pos = chunk_text.rfind(char)
                if pos > best_break and pos > self._chunk_size // 2:
                    best_break = pos + len(char)

            if best_break > 0:
                end = start + best_break
            else:
                # Fall back to word boundary
                space_pos = chunk_text.rfind(" ")
                if space_pos > self._chunk_size // 2:
                    end = start + space_pos + 1

            chunks.append(text[start:end].strip())

            # Move start with overlap
            start = end - self._chunk_overlap
            if start < 0:
                start = 0

        return [c for c in chunks if c]

    @staticmethod
    def _generate_doc_id(path: Path, chunk_index: int) -> str:
        """Generate a unique document ID based on file path and chunk index.

        Args:
            path: File path.
            chunk_index: Index of the chunk.

        Returns:
            Unique document ID.
        """
        content = f"{path.absolute()}:{chunk_index}"
        return hashlib.sha256(content.encode()).hexdigest()[:16]


def create_chunker(
    chunk_size: int = 512,
    chunk_overlap: int = 50,
    pdf_strategy: Literal["fast", "hi_res", "ocr_only"] = "fast",
) -> DocumentChunker:
    """Factory function to create a document chunker.

    Args:
        chunk_size: Target size of chunks in characters.
        chunk_overlap: Number of characters to overlap between chunks.
        pdf_strategy: Strategy for PDF processing.

    Returns:
        Configured DocumentChunker instance.
    """
    return DocumentChunker(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        pdf_strategy=pdf_strategy,
    )

